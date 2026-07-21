"""
Document chunking service.

Processes benchmark datasets or uploaded files (PDF, DOCX, TXT) using LangChain's
RecursiveCharacterTextSplitter. Generates structured chunk data with stable,
deterministic identifiers and metadata.
"""

from __future__ import annotations

import logging
import uuid
from datetime import datetime, timezone
from pathlib import Path
from typing import Any, List

from fastapi import HTTPException
import docx
from pypdf import PdfReader

from schemas.chunk import Chunk, ChunkMetadata
from services.dataset_service import get_truthfulqa_records, get_squad_records

logger = logging.getLogger("llm_evaluation")


class ChunkingService:
    """Service to extract text from multiple formats and split it into semantic chunks."""

    def __init__(self, chunk_size: int = 500, chunk_overlap: int = 100) -> None:
        """Initialise the service with standard character splitter parameters."""
        try:
            from langchain.text_splitter import RecursiveCharacterTextSplitter
        except ImportError:
            from langchain_text_splitters import RecursiveCharacterTextSplitter  # type: ignore[import-not-found]

        self.text_splitter = RecursiveCharacterTextSplitter(
            chunk_size=chunk_size,
            chunk_overlap=chunk_overlap,
            separators=["\n\n", "\n", ". ", " ", ""],
        )

    def _extract_text_from_pdf(self, file_path: Path) -> str:
        """Extract text from a PDF document page by page."""
        logger.info("[CHUNK] Extracting text from PDF: %s", file_path.name)
        file_size_mb = file_path.stat().st_size / (1024 * 1024)
        logger.info("[CHUNK] PDF file size: %.2f MB", file_size_mb)

        if file_size_mb > 10.0:
            logger.warning("[CHUNK] Oversized PDF file detected: %.2f MB", file_size_mb)

        try:
            reader = PdfReader(file_path)
            num_pages = len(reader.pages)
            logger.info("[CHUNK] PDF contains %d pages", num_pages)

            text_parts = []
            for i in range(num_pages):
                try:
                    page = reader.pages[i]
                    page_text = page.extract_text()
                    if page_text:
                        text_parts.append(page_text)
                except Exception as page_exc:  # noqa: BLE001
                    logger.warning(
                        "[CHUNK] Page %d extraction skipped in %s: %s",
                        i, file_path.name, page_exc
                    )

            extracted_text = "\n".join(text_parts).strip()
            return extracted_text
        except Exception as exc:
            logger.error("[CHUNK] Error reading PDF %s: %s", file_path.name, exc)
            raise ValueError(f"Failed to parse PDF file '{file_path.name}': {exc}") from exc

    def _extract_text_from_docx(self, file_path: Path) -> str:
        """Extract text from a Word document (.docx)."""
        logger.info("[CHUNK] Extracting text from DOCX: %s", file_path.name)
        file_size_mb = file_path.stat().st_size / (1024 * 1024)
        if file_size_mb > 10.0:
            logger.warning("[CHUNK] Oversized DOCX file detected: %.2f MB", file_size_mb)

        try:
            doc = docx.Document(str(file_path))
            paragraphs = [p.text for p in doc.paragraphs if p.text is not None]

            # Also extract tables content for complete document parsing
            table_text = []
            for table in doc.tables:
                for row in table.rows:
                    row_text = [cell.text for cell in row.cells if cell.text]
                    if row_text:
                        table_text.append(" | ".join(row_text))

            all_text = paragraphs + table_text
            extracted_text = "\n".join(all_text).strip()
            return extracted_text
        except Exception as exc:
            logger.error("[CHUNK] Error reading DOCX %s: %s", file_path.name, exc)
            raise ValueError(f"Failed to parse Word document '{file_path.name}': {exc}") from exc

    def _extract_text_from_txt(self, file_path: Path) -> str:
        """Extract text from a plain text file, attempting common encodings."""
        logger.info("[CHUNK] Extracting text from TXT: %s", file_path.name)
        encodings = ["utf-8", "latin-1", "cp1252"]
        
        for encoding in encodings:
            try:
                with open(file_path, "r", encoding=encoding) as f:
                    extracted_text = f.read().strip()
                return extracted_text
            except UnicodeDecodeError:
                continue

        raise ValueError(
            f"Failed to decode text file '{file_path.name}' with standard encodings."
        )

    def extract_text_from_file(self, file_path: Path) -> str:
        """Determine file type and extract its text content."""
        if not file_path.exists():
            raise FileNotFoundError(f"File not found: {file_path}")

        ext = file_path.suffix.lower()
        if ext == ".pdf":
            return self._extract_text_from_pdf(file_path)
        elif ext == ".docx":
            return self._extract_text_from_docx(file_path)
        elif ext == ".txt":
            return self._extract_text_from_txt(file_path)
        else:
            raise ValueError(
                f"Unsupported file type '{ext}'. Allowed types: .pdf, .docx, .txt"
            )

    def chunk_text(
        self,
        text: str,
        source: str,
        document_name: str,
        document_id: uuid.UUID,
    ) -> list[Chunk]:
        """Split raw text into structured Chunk models using RecursiveCharacterTextSplitter."""
        # Reject empty text
        if not text or not text.strip():
            logger.warning("[CHUNK] Rejected empty/whitespace text for %s", document_name)
            raise ValueError("Document text is empty or contains only whitespace.")

        splits = self.text_splitter.split_text(text)
        chunks: list[Chunk] = []
        created_at = datetime.now(timezone.utc).isoformat().replace("+00:00", "Z")

        for idx, split in enumerate(splits):
            chunk_text = split.strip()
            if not chunk_text:
                continue

            metadata = ChunkMetadata(
                document_name=document_name,
                created_at=created_at,
                chunk_size=len(chunk_text),
            )
            chunk = Chunk(
                chunk_id=uuid.uuid5(document_id, f"chunk-{idx}"),
                document_id=document_id,
                chunk_index=idx,
                text=chunk_text,
                source=source,
                metadata=metadata,
            )
            chunks.append(chunk)

        return chunks

    def process_chunking(self, source: str) -> list[Chunk]:
        """Orchestrate chunking process based on the selected input source."""
        logger.info("[CHUNK] Starting chunking run for source: %s", source)
        all_chunks: list[Chunk] = []

        if source in ("truthfulqa", "squad"):
            # Load dataset records dynamically
            if source == "truthfulqa":
                records = get_truthfulqa_records()
            else:
                records = get_squad_records()

            logger.info("[CHUNK] Loaded %d records from dataset: %s", len(records), source)

            for record in records:
                # Format text representation of record
                record_text = f"Question: {record.question}\nAnswer: {record.answer}"
                
                # Generate deterministic UUID based on namespace & unique record ID
                doc_id = uuid.uuid5(uuid.NAMESPACE_DNS, f"{source}-{record.id}")
                doc_name = f"{source}_{record.id}"

                try:
                    record_chunks = self.chunk_text(
                        text=record_text,
                        source=source,
                        document_name=doc_name,
                        document_id=doc_id,
                    )
                    all_chunks.extend(record_chunks)
                except ValueError as val_err:
                    # Skip empty record text and log at debug level
                    logger.debug("[CHUNK] Skipped record %s: %s", record.id, val_err)
                    continue

        elif source == "uploaded_document":
            uploads_dir = Path(__file__).resolve().parent.parent / "uploads"
            if not uploads_dir.exists():
                uploads_dir.mkdir(parents=True, exist_ok=True)

            files = [
                f for f in uploads_dir.iterdir()
                if f.is_file() and f.name != ".gitkeep"
            ]

            if not files:
                logger.warning("[CHUNK] Uploads folder is empty")
                raise HTTPException(
                    status_code=400,
                    detail=(
                        "No uploaded documents found. Please upload PDF, DOCX, "
                        "or TXT documents first."
                    ),
                )

            logger.info("[CHUNK] Found %d files in uploads directory", len(files))

            for file_path in files:
                try:
                    # Generate deterministic document ID based on filename
                    doc_id = uuid.uuid5(uuid.NAMESPACE_DNS, f"file-{file_path.name}")
                    doc_name = file_path.name

                    # Extract text content
                    text = self.extract_text_from_file(file_path)

                    # Generate chunks
                    doc_chunks = self.chunk_text(
                        text=text,
                        source=source,
                        document_name=doc_name,
                        document_id=doc_id,
                    )
                    all_chunks.extend(doc_chunks)
                except ValueError as val_err:
                    logger.warning(
                        "[CHUNK] Skipping file %s due to parsing/validation failure: %s",
                        file_path.name, val_err
                    )
                    continue
                except Exception as exc:  # noqa: BLE001
                    logger.error("[CHUNK] Failed to process file %s: %s", file_path.name, exc)
                    continue

            if not all_chunks:
                raise HTTPException(
                    status_code=422,
                    detail="All uploaded files were either empty or could not be parsed successfully.",
                )

        else:
            logger.error("[CHUNK] Unsupported chunk source requested: %s", source)
            raise HTTPException(
                status_code=422,
                detail=f"Unsupported source '{source}'. Valid sources are: truthfulqa, squad, uploaded_document",
            )

        logger.info("[CHUNK] Successfully generated %d chunks from %s", len(all_chunks), source)
        return all_chunks
